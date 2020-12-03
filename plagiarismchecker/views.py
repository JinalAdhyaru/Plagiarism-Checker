from django.shortcuts import render
from django.http import HttpResponse
from plagiarismchecker.algorithm import main
from docx import *

from plagiarismchecker.algorithm import fileSimilarity
#import matplotlib.pyplot as plt
#from plagiarismchecker.algorithm import main

# Create your views here.
def home(request):
    return render(request, 'pc/index.html') 

def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text
        
    percent,link = main.findSimilarity(value)
    percent = round(percent,2)
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    
#two file compare
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())
        
    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    result = round(result,2)
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})


